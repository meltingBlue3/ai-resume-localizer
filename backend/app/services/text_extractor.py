import io

import pymupdf
from docx import Document


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF with reading-order sort."""
    doc = pymupdf.open(stream=content, filetype="pdf")
    pages: list[str] = []
    try:
        for page in doc:
            pages.append(page.get_text("text", sort=True))
    finally:
        doc.close()
    return "\n\n".join(pages).strip()


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes, including table cell content.

    Chinese resumes often use table layouts, so we iterate both
    paragraphs and tables to capture all text.
    """
    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()
